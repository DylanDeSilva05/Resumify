"""
CV Parsing service for extracting structured information from CV files
"""
import os
import re
import json
import logging
from typing import Dict, List, Any, Optional, Tuple
from pathlib import Path
import PyPDF2
from docx import Document
import spacy
from sklearn.feature_extraction.text import TfidfVectorizer
import pandas as pd
from datetime import datetime
import email_validator

from app.core.config import settings

logger = logging.getLogger(__name__)


class CVParser:
    """Service class for parsing CV files and extracting structured information"""

    def __init__(self):
        """Initialize CV parser with NLP model"""
        try:
            self.nlp = spacy.load(settings.SPACY_MODEL)
        except OSError:
            logger.warning(f"SpaCy model '{settings.SPACY_MODEL}' not found. Using basic parsing.")
            self.nlp = None

        # Import the comprehensive skills database from NLP service
        from app.services.nlp_service import NLPService
        nlp_service = NLPService()

        # Use the same comprehensive skill databases as NLP service
        self.technical_skills = nlp_service.technical_skills_db
        self.soft_skills = nlp_service.soft_skills_db

        # Flatten technical skills for easier searching
        self.all_technical_skills = nlp_service.all_technical_skills

        # Education keywords
        self.education_keywords = [
            'bachelor', 'master', 'phd', 'doctorate', 'degree', 'university', 'college',
            'institute', 'school', 'education', 'graduated', 'gpa'
        ]

        # Experience keywords
        self.experience_keywords = [
            'work', 'experience', 'employment', 'position', 'job', 'role', 'company',
            'organization', 'years', 'months', 'present', 'current'
        ]

    def parse_cv_file(self, file_path: str) -> Dict[str, Any]:
        """
        Parse CV file and extract structured information

        Args:
            file_path: Path to the CV file

        Returns:
            Dict: Structured CV information
        """
        try:
            # Extract text from file
            raw_text = self._extract_text_from_file(file_path)
            if not raw_text:
                raise ValueError("Could not extract text from file")

            # Parse structured information
            parsed_data = self._parse_text_content(raw_text)
            parsed_data['raw_text'] = raw_text
            parsed_data['parsing_status'] = 'completed'

            logger.info(f"Successfully parsed CV: {file_path}")
            return parsed_data

        except Exception as e:
            logger.error(f"Failed to parse CV {file_path}: {str(e)}")
            return {
                'raw_text': '',
                'parsing_status': 'failed',
                'parsing_error': str(e),
                'personal_info': {},
                'education': [],
                'work_experience': [],
                'skills': {'technical': [], 'soft': []},
                'certifications': [],
                'languages': [],
                'total_experience_years': 0.0
            }

    def _extract_text_from_file(self, file_path: str) -> str:
        """Extract raw text from CV file"""
        file_extension = Path(file_path).suffix.lower()

        if file_extension == '.pdf':
            return self._extract_from_pdf(file_path)
        elif file_extension in ['.doc', '.docx']:
            return self._extract_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")

    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
        except Exception as e:
            logger.error(f"Failed to extract text from PDF {file_path}: {str(e)}")
            raise
        return text.strip()

    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file including tables"""
        try:
            doc = Document(file_path)
            text_parts = []

            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        # Join cells with a space and add to text
                        text_parts.append(' '.join(row_text))

            text = "\n".join(text_parts)
        except Exception as e:
            logger.error(f"Failed to extract text from DOCX {file_path}: {str(e)}")
            raise
        return text.strip()

    def _parse_text_content(self, text: str) -> Dict[str, Any]:
        """Parse raw text and extract structured information"""
        lines = [line.strip() for line in text.split('\n') if line.strip()]

        return {
            'personal_info': self._extract_personal_info(text),
            'education': self._extract_education(text),
            'work_experience': self._extract_work_experience(text, lines),
            'skills': self._extract_skills(text),
            'certifications': self._extract_certifications(text),
            'languages': self._extract_languages(text),
            'total_experience_years': self._calculate_experience_years(text)
        }

    def _extract_personal_info(self, text: str) -> Dict[str, Any]:
        """Extract personal information from CV text"""
        info = {}

        # Extract email
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, text)
        if emails:
            info['email'] = emails[0]

        # Extract phone number
        phone_patterns = [
            r'\b\d{3}[-.]?\d{3}[-.]?\d{4}\b',
            r'\(\d{3}\)\s*\d{3}[-.]?\d{4}',
            r'\+\d{1,3}[-.\s]?\d{1,14}',
        ]
        for pattern in phone_patterns:
            phones = re.findall(pattern, text)
            if phones:
                info['phone'] = phones[0]
                break

        # Extract location (basic implementation)
        location_keywords = ['city', 'state', 'country', 'address']
        for keyword in location_keywords:
            if keyword in text.lower():
                # This is a simplified implementation
                # In production, you'd use more sophisticated NLP
                info['location'] = "Location extracted"  # Placeholder
                break

        return info

    def _extract_education(self, text: str) -> List[Dict[str, Any]]:
        """Extract education information from CV text"""
        education = []
        lines = [line.strip() for line in text.split('\n') if line.strip()]

        # Find the education section
        education_section_start = -1
        education_section_end = len(lines)

        # Keywords that indicate start of education section
        education_headers = [
            'education', 'academic background', 'academic qualifications',
            'educational background', 'qualifications', 'academic history',
            'degrees', 'education and training'
        ]

        # Keywords that indicate end of education section
        next_section_headers = [
            'experience', 'work', 'employment', 'skills', 'certifications',
            'projects', 'awards', 'references', 'publications', 'languages',
            'interests', 'hobbies', 'professional experience'
        ]

        # Find education section boundaries
        for i, line in enumerate(lines):
            line_lower = line.lower().strip()

            # Check if this is the start of education section
            if education_section_start == -1:
                for header in education_headers:
                    if line_lower == header or line_lower.startswith(header + ':'):
                        education_section_start = i + 1
                        break
            # Check if we've reached the next section
            elif education_section_start != -1:
                for header in next_section_headers:
                    if line_lower == header or line_lower.startswith(header + ':'):
                        education_section_end = i
                        break
                if education_section_end != len(lines):
                    break

        # If no education section found, search entire text
        if education_section_start == -1:
            education_section_start = 0

        section_lines = lines[education_section_start:education_section_end]

        # Degree patterns and keywords
        degree_keywords = {
            'bachelor': ['bachelor', 'b.s.', 'b.a.', 'b.sc.', 'b.tech', 'b.e.', 'bs', 'ba', 'bsc', 'btech'],
            'master': ['master', 'm.s.', 'm.a.', 'm.sc.', 'm.tech', 'm.e.', 'ms', 'ma', 'msc', 'mtech', 'mba'],
            'phd': ['phd', 'ph.d.', 'doctorate', 'doctoral', 'd.phil'],
            'diploma': ['diploma', 'certificate', 'associate'],
        }

        # Institution indicators
        institution_keywords = [
            'university', 'college', 'institute', 'school', 'academy',
            'polytechnic', 'conservatory', 'seminary'
        ]

        # Use spaCy NER if available to identify organizations (universities)
        organizations = []
        if self.nlp:
            try:
                section_text = '\n'.join(section_lines)
                doc = self.nlp(section_text)
                organizations = [ent.text for ent in doc.ents if ent.label_ == "ORG"]
            except Exception as e:
                logger.warning(f"NER organization extraction for education failed: {str(e)}")

        # Parse education entries
        i = 0
        while i < len(section_lines):
            line = section_lines[i].strip()
            if not line:
                i += 1
                continue

            line_lower = line.lower()
            degree = None
            institution = None
            year = None

            # Check if line contains degree keywords
            degree_found = False
            for degree_type, patterns in degree_keywords.items():
                for pattern in patterns:
                    # Use word boundary to avoid partial matches
                    if re.search(rf'\b{re.escape(pattern)}\b', line_lower):
                        degree = line  # Use the full line as degree initially
                        degree_found = True
                        break
                if degree_found:
                    break

            # Look for institution in current and nearby lines
            if degree_found or any(keyword in line_lower for keyword in institution_keywords):
                # Extract year from nearby lines
                search_range = min(i + 3, len(section_lines))
                for j in range(i, search_range):
                    check_line = section_lines[j]

                    # Look for year (graduation year)
                    year_matches = re.findall(r'\b(19|20)\d{2}\b', check_line)
                    if year_matches and not year:
                        # Take the last year mentioned (usually graduation year)
                        year = int(year_matches[-1])

                    # Look for institution
                    if not institution:
                        check_lower = check_line.lower()
                        # Check if line contains institution keywords
                        if any(keyword in check_lower for keyword in institution_keywords):
                            institution = check_line.strip()

                        # Check against NER organizations
                        elif organizations:
                            for org in organizations:
                                if org.lower() in check_lower:
                                    institution = org
                                    break

                # If we didn't find a clear degree, try to extract it from the pattern
                if degree and not degree_found:
                    # This line might be institution, not degree
                    if any(keyword in line_lower for keyword in institution_keywords):
                        institution = line
                        degree = None
                    else:
                        # Try to extract degree from context
                        for degree_type, patterns in degree_keywords.items():
                            for pattern in patterns:
                                if re.search(rf'\b{re.escape(pattern)}\b', line_lower):
                                    # Extract the degree part
                                    degree_match = re.search(rf'.*?\b{re.escape(pattern)}\b.*', line, re.IGNORECASE)
                                    if degree_match:
                                        degree = degree_match.group(0).strip()
                                    break

                # If we found institution but no degree, look backwards for degree
                if institution and not degree:
                    for j in range(max(0, i-2), i):
                        check_line = section_lines[j].lower()
                        for degree_type, patterns in degree_keywords.items():
                            for pattern in patterns:
                                if re.search(rf'\b{re.escape(pattern)}\b', check_line):
                                    degree = section_lines[j].strip()
                                    break
                            if degree:
                                break
                        if degree:
                            break

                # Add to education list if we have either degree or institution
                if degree or institution:
                    # Clean up degree and institution
                    if degree and institution and degree == institution:
                        # If they're the same, one is likely wrong
                        if any(keyword in degree.lower() for keyword in institution_keywords):
                            institution = degree
                            degree = None
                        else:
                            institution = None

                    education.append({
                        'degree': degree if degree else 'Degree not specified',
                        'institution': institution if institution else 'Institution not specified',
                        'year': year
                    })

                # Skip processed lines
                i = min(search_range, i + 1)
            else:
                i += 1

        # Remove duplicates based on degree and institution
        seen = set()
        unique_education = []
        for edu in education:
            key = (edu['degree'].lower(), edu['institution'].lower() if edu['institution'] else '')
            if key not in seen:
                seen.add(key)
                unique_education.append(edu)

        return unique_education

    def _extract_work_experience(self, text: str, lines: List[str]) -> List[Dict[str, Any]]:
        """Extract work experience from CV text"""
        experience = []

        # Find the work experience section
        experience_section_start = -1
        experience_section_end = len(lines)

        # Keywords that indicate start of experience section
        experience_headers = [
            'work experience', 'professional experience', 'employment history',
            'work history', 'experience', 'employment', 'professional background',
            'career history', 'relevant experience'
        ]

        # Keywords that indicate end of experience section (start of another section)
        next_section_headers = [
            'education', 'skills', 'certifications', 'projects', 'awards',
            'references', 'publications', 'languages', 'interests', 'hobbies'
        ]

        # Find experience section boundaries
        for i, line in enumerate(lines):
            line_lower = line.lower().strip()

            # Check if this is the start of experience section
            if experience_section_start == -1:
                for header in experience_headers:
                    if line_lower == header or line_lower.startswith(header + ':'):
                        experience_section_start = i + 1
                        break
            # Check if we've reached the next section
            elif experience_section_start != -1:
                for header in next_section_headers:
                    if line_lower == header or line_lower.startswith(header + ':'):
                        experience_section_end = i
                        break
                if experience_section_end != len(lines):
                    break

        # If no experience section found, try to extract from entire text
        if experience_section_start == -1:
            experience_section_start = 0

        # Extract experience entries from the section
        section_lines = lines[experience_section_start:experience_section_end]

        if not section_lines:
            return experience

        # Use spaCy NER if available to identify organizations
        organizations = []
        if self.nlp:
            try:
                section_text = '\n'.join(section_lines)
                doc = self.nlp(section_text)
                organizations = [ent.text for ent in doc.ents if ent.label_ == "ORG"]
            except Exception as e:
                logger.warning(f"NER organization extraction failed: {str(e)}")

        # Common job title patterns and keywords
        job_title_keywords = [
            'engineer', 'developer', 'manager', 'analyst', 'scientist', 'designer',
            'consultant', 'specialist', 'coordinator', 'director', 'assistant',
            'officer', 'lead', 'senior', 'junior', 'associate', 'administrator',
            'technician', 'supervisor', 'executive', 'architect', 'researcher',
            'professor', 'teacher', 'instructor', 'doctor', 'nurse', 'accountant',
            'programmer', 'tester', 'qa', 'devops', 'sre', 'intern', 'trainee'
        ]

        # Company suffix patterns
        company_patterns = [
            r'\b(inc\.?|ltd\.?|llc|corp\.?|corporation|company|co\.?|gmbh|pvt\.?|limited)\b',
            r'\b(technologies|systems|solutions|services|consulting|group)\b'
        ]

        # Parse experience entries
        i = 0
        while i < len(section_lines):
            line = section_lines[i].strip()
            if not line:
                i += 1
                continue

            # Look for job title (usually contains job keywords)
            potential_title = None
            line_lower = line.lower()

            # Check if line contains job title keywords
            if any(keyword in line_lower for keyword in job_title_keywords):
                potential_title = line

            # Look for date range patterns in this line or next few lines
            date_pattern = r'\b(19|20)\d{2}\b.*?(?:[-–—]|to)\s*(?:\b(19|20)\d{2}\b|present|current)'
            dates_in_line = re.search(date_pattern, line, re.IGNORECASE)

            # If we have a potential title, look for company and dates nearby
            if potential_title or dates_in_line:
                position = potential_title if potential_title else line
                company = None
                start_date = None
                end_date = None
                description_lines = []

                # Extract dates from current or nearby lines
                search_range = min(i + 3, len(section_lines))
                for j in range(i, search_range):
                    dates_match = re.search(date_pattern, section_lines[j], re.IGNORECASE)
                    if dates_match:
                        # Extract start year
                        start_match = re.search(r'\b(19|20)\d{2}\b', dates_match.group(0))
                        if start_match:
                            start_date = int(start_match.group(0))

                        # Extract end year or set to current year if "present"
                        if 'present' in dates_match.group(0).lower() or 'current' in dates_match.group(0).lower():
                            end_date = datetime.now().year
                        else:
                            end_years = re.findall(r'\b(19|20)\d{2}\b', dates_match.group(0))
                            if len(end_years) > 1:
                                end_date = int(end_years[-1])
                            elif len(end_years) == 1:
                                end_date = int(end_years[0])
                        break

                # Look for company name (check nearby lines and NER results)
                for j in range(i, min(i + 3, len(section_lines))):
                    check_line = section_lines[j]
                    # Check if line matches company patterns
                    for pattern in company_patterns:
                        if re.search(pattern, check_line, re.IGNORECASE):
                            company = check_line.strip()
                            break

                    # Check against NER organizations
                    if not company and organizations:
                        for org in organizations:
                            if org.lower() in check_line.lower():
                                company = org
                                break

                    if company:
                        break

                # Extract description (bullet points or text after the header info)
                desc_start = i + 1
                if dates_in_line:
                    desc_start = i + 1
                else:
                    # Skip lines that are part of header (title, company, dates)
                    while desc_start < min(i + 3, len(section_lines)):
                        if re.search(date_pattern, section_lines[desc_start], re.IGNORECASE):
                            desc_start += 1
                            break
                        desc_start += 1

                # Collect description lines until next job entry or end
                for j in range(desc_start, min(desc_start + 10, len(section_lines))):
                    desc_line = section_lines[j].strip()

                    # Stop if we hit what looks like another job title
                    if any(keyword in desc_line.lower() for keyword in job_title_keywords):
                        # But only if it's not a bullet point continuing the description
                        if not desc_line.startswith(('•', '-', '*', '–')):
                            break

                    # Stop if we hit a date range (likely next job)
                    if re.search(date_pattern, desc_line, re.IGNORECASE):
                        break

                    if desc_line:
                        description_lines.append(desc_line)

                # Calculate duration
                duration_months = 0
                if start_date and end_date:
                    duration_months = max(1, (end_date - start_date) * 12)

                # Add to experience list
                if position or company:  # At least one should be present
                    experience.append({
                        'position': position if position else 'Not specified',
                        'company': company if company else 'Not specified',
                        'start_date': start_date,
                        'end_date': end_date,
                        'description': ' '.join(description_lines[:3]) if description_lines else '',  # First 3 lines
                        'duration_months': duration_months
                    })

                # Skip processed lines
                i = desc_start + len(description_lines)
            else:
                i += 1

        return experience

    def _extract_skills(self, text: str) -> Dict[str, List[str]]:
        """Extract technical and soft skills from CV text"""
        text_lower = text.lower()

        # Extract technical skills
        technical_skills = []
        for skill in self.all_technical_skills:
            # Use word boundaries to avoid partial matches
            pattern = rf'\b{re.escape(skill.lower())}\b'
            if re.search(pattern, text_lower):
                technical_skills.append(skill)

        # Extract soft skills
        soft_skills = []
        for skill in self.soft_skills:
            # Handle multi-word skills properly
            pattern = rf'\b{re.escape(skill.lower())}\b'
            if re.search(pattern, text_lower):
                soft_skills.append(skill)

        logger.info(f"Extracted {len(technical_skills)} technical skills: {technical_skills}")
        logger.info(f"Extracted {len(soft_skills)} soft skills: {soft_skills}")

        return {
            'technical': list(set(technical_skills)),
            'soft': list(set(soft_skills))
        }

    def _extract_certifications(self, text: str) -> List[Dict[str, Any]]:
        """Extract certifications from CV text"""
        certifications = []
        cert_keywords = ['certificate', 'certification', 'certified', 'license']

        for keyword in cert_keywords:
            if keyword.lower() in text.lower():
                certifications.append({
                    'name': 'Certification extracted',
                    'issuer': 'Issuer extracted',
                    'year': self._extract_years_from_context(text, keyword)
                })

        return certifications

    def _extract_languages(self, text: str) -> List[Dict[str, str]]:
        """Extract language skills from CV text"""
        languages = []
        common_languages = [
            'english', 'spanish', 'french', 'german', 'chinese', 'japanese',
            'portuguese', 'russian', 'arabic', 'hindi', 'italian'
        ]

        for language in common_languages:
            if language.lower() in text.lower():
                languages.append({
                    'language': language.title(),
                    'proficiency': 'Native/Fluent'  # Simplified
                })

        return languages

    def _calculate_experience_years(self, text: str) -> float:
        """Calculate total years of experience from CV text"""
        # Look for experience patterns
        experience_patterns = [
            r'(\d+)\s*years?\s*(of\s*)?experience',
            r'experience\s*:\s*(\d+)\s*years?',
            r'(\d+)\+?\s*years?\s*experience'
        ]

        years = []
        for pattern in experience_patterns:
            matches = re.findall(pattern, text.lower())
            for match in matches:
                try:
                    years.append(float(match[0] if isinstance(match, tuple) else match))
                except (ValueError, IndexError):
                    continue

        if years:
            return max(years)  # Return the highest experience mentioned

        # Fallback: calculate from work experience dates
        all_years = self._extract_all_years(text)
        if len(all_years) >= 2:
            return float(max(all_years) - min(all_years))

        return 0.0

    def _extract_years_from_context(self, text: str, context: str) -> Optional[int]:
        """Extract year from text context"""
        # Look for 4-digit years near the context
        context_index = text.lower().find(context.lower())
        if context_index != -1:
            context_area = text[max(0, context_index-50):context_index+50]
            years = re.findall(r'\b(19|20)\d{2}\b', context_area)
            if years:
                return int(years[0])
        return None

    def _extract_all_years(self, text: str) -> List[int]:
        """Extract all 4-digit years from text"""
        years = re.findall(r'\b(19|20)\d{2}\b', text)
        return [int(year) for year in years]

    def extract_candidate_name(self, text: str) -> str:
        """Extract candidate name from CV text using NLP and pattern matching"""
        if not text or not text.strip():
            return "Unknown"

        # Strategy 1: Use spaCy NER to find PERSON entities
        if self.nlp:
            try:
                # Process the first 500 characters where name is likely to appear
                doc = self.nlp(text[:500])
                person_entities = [ent.text for ent in doc.ents if ent.label_ == "PERSON"]

                if person_entities:
                    # Clean and validate the first person entity found
                    name = person_entities[0].strip()

                    # Remove trailing keywords (email, phone, etc.) from NER result
                    section_keywords = ['email', 'phone', 'address', 'name', 'tel', 'mobile', 'fax', 'candidate', 'contact', 'dob', 'date']
                    for keyword in section_keywords:
                        # Remove keyword with any preceding whitespace (including newlines)
                        pattern = rf'\s+{re.escape(keyword)}$'
                        name = re.sub(pattern, '', name, flags=re.IGNORECASE).strip()

                    # Validate it's a reasonable name (1-5 words, mostly letters, no invalid keywords)
                    name_words = name.split()
                    if 1 <= len(name_words) <= 5 and len(name) >= 2:
                        # Check that words don't contain CV section keywords
                        invalid_words = ['email', 'phone', 'address', 'contact', 'information', 'details']
                        if not any(word.lower() in invalid_words for word in name_words):
                            return name.title()
            except Exception as e:
                logger.warning(f"NER name extraction failed: {str(e)}")

        # Strategy 2: Check first few non-empty lines
        lines = [line.strip() for line in text.split('\n') if line.strip()]

        if not lines:
            return "Unknown"

        # Common CV headers to skip
        headers_to_skip = [
            'curriculum vitae', 'resume', 'cv', 'professional resume',
            'personal resume', 'contact', 'contact information',
            'objective', 'summary', 'professional summary',
            'personal information', 'personal details', 'address',
            'profile', 'about me', 'career objective', 'career summary'
        ]

        # Check first 10 lines for a valid name
        for i, line in enumerate(lines[:10]):
            line_lower = line.lower().strip()

            # Skip header lines (lines that ARE section headers)
            if any(header == line_lower for header in headers_to_skip):
                continue

            # Skip lines that are ONLY section keywords (but allow "Name: John Doe" format)
            skip_if_only = ['email', 'phone', 'address', 'linkedin', 'information', 'details', 'profile', 'education', 'experience', 'skills', 'certification']
            if line_lower in skip_if_only:
                continue

            # Skip lines with @ (email addresses), URLs
            if '@' in line or 'http' in line_lower or 'www.' in line_lower or '.com' in line_lower:
                continue

            # Check if line looks like a name
            # Handle "Name: John Doe" or "Candidate Name: John Doe" format
            if ':' in line:
                parts = line.split(':', 1)
                label = parts[0].strip().lower()
                value = parts[1].strip() if len(parts) > 1 else ""

                # If the label indicates this is a name field, extract the value
                if 'name' in label and value and not any(word in label for word in ['company', 'file', 'user']):
                    cleaned_line = value
                else:
                    # Not a name field, skip this line
                    continue
            else:
                # No colon, process as regular line
                cleaned_line = line

            # Remove common punctuation but keep spaces
            cleaned_line = re.sub(r'[^\w\s\-\.]', ' ', cleaned_line).strip()

            # Remove any standalone CV section words
            tokens = cleaned_line.split()
            filtered_tokens = []

            # Extended list of CV section keywords to filter out
            section_keywords = [
                'email', 'phone', 'address', 'name', 'tel', 'mobile', 'fax', 'candidate',
                'telephone', 'contact', 'cell', 'gmail', 'yahoo', 'hotmail', 'outlook'
            ]

            for token in tokens:
                token_lower = token.lower()
                # Skip common standalone CV section words
                if token_lower not in section_keywords:
                    filtered_tokens.append(token)

            if not filtered_tokens:
                continue

            cleaned_line = ' '.join(filtered_tokens)

            # Additional cleanup: Remove trailing CV keywords using regex
            # This catches cases like "John Doe Email", "John Doe\nEmail", etc.
            trailing_keywords = ['email', 'phone', 'address', 'contact', 'tel', 'mobile', 'name']
            for keyword in trailing_keywords:
                # Use regex to match keyword at end with any whitespace before it
                pattern = rf'\s+{re.escape(keyword)}$'
                cleaned_line = re.sub(pattern, '', cleaned_line, flags=re.IGNORECASE).strip()

                # Also check if the whole line is just the keyword
                if cleaned_line.lower() == keyword:
                    cleaned_line = ''
                    break

            if not cleaned_line:
                continue

            words = cleaned_line.split()

            # Valid name criteria:
            # - 1-5 words (after filtering)
            # - Each word is mostly alphabetic
            # - Total length is reasonable (2-50 characters)
            # - Not all uppercase (unless 2 words or less)
            if 1 <= len(words) <= 5 and 2 <= len(cleaned_line) <= 50:
                # Check if words are mostly alphabetic
                mostly_alpha = all(
                    sum(c.isalpha() for c in word) / len(word) >= 0.7
                    for word in words if len(word) > 0
                )

                if mostly_alpha:
                    # Avoid all-caps names unless they're short
                    if len(words) <= 2 or not cleaned_line.isupper():
                        # Final validation: ensure name doesn't contain CV section words
                        name_lower = cleaned_line.lower()
                        invalid_words = ['information', 'details', 'address', 'personal', 'contact',
                                       'profile', 'objective', 'summary', 'education', 'experience']
                        if not any(word in name_lower for word in invalid_words):
                            return cleaned_line.title()

        # Strategy 3: Extract from email if available
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, text[:1000])
        if emails:
            # Extract name from email (e.g., john.doe@example.com -> John Doe)
            email_username = emails[0].split('@')[0]
            # Replace dots, underscores, and digits with spaces
            name_from_email = re.sub(r'[._\d]+', ' ', email_username).strip()
            if name_from_email and len(name_from_email) >= 2:
                return name_from_email.title()

        return "Unknown"